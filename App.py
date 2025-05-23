import streamlit as st
from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO


def aplicar_highlight(paragraph):
    """Aplica fundo amarelo (marca-texto) no parágrafo inteiro."""
    for run in paragraph.runs:
        rPr = run._element.get_or_add_rPr()
        highlight = OxmlElement('w:highlight')
        highlight.set(qn('w:val'), 'yellow')
        rPr.append(highlight)


def gerar_legendas(publicacao):
    texto = publicacao.lower()
    legendas = []

    if "acórdão" in texto:
        legendas.append("ANALISAR ACÓRDÃO")
        legendas.append("EMAIL ACÓRDÃO")
        legendas.append("LEMBRETE RECURSAL")
    elif "sentença homologatória" in texto or "homologo" in texto:
        legendas.append("ANALISAR SENTENÇA HOMOLOGATÓRIA")
        legendas.append("EMAIL SENTENÇA HOMOLOGATÓRIA")
        legendas.append("PAGAMENTO CLIENTE")
    elif "sentença" in texto:
        legendas.append("ANALISAR SENTENÇA")
        legendas.append("EMAIL SENTENÇA")
    elif "despacho" in texto and "deneg" in texto:
        legendas.append("AGRAVOIRREV")
    elif "despacho" in texto:
        legendas.append("ANALISAR DESPACHO")
    elif "perícia" in texto or "laudo" in texto:
        legendas.append("MANIFESTAÇÃO DE LAUDO")
        legendas.append("EMAIL LAUDO")
    elif "redesignação de perícia" in texto:
        legendas.append("AGENDAR PERÍCIA")
        legendas.append("EMAIL PERÍCIA")
    elif "pauta" in texto or "audiência" in texto:
        legendas.append("AGENDAR AUDIÊNCIA")
    elif "alvará" in texto or "liberação de valores" in texto or "crédito" in texto or "expedição" in texto:
        legendas.append("PAGAMENTO CLIENTE")
        legendas.append("EMAIL ROSY PAGAMENTO")
    elif "agravo" in texto:
        if "petição" in texto:
            legendas.append("AGRAVO DE PETIÇÃO")
        else:
            legendas.append("AGRAVO DE INSTRUMENTO")
    elif "recurso de revista" in texto or "rr" in texto:
        legendas.append("RR")
    elif "embargos de declaração" in texto or "ed" in texto:
        legendas.append("ED")
    elif "execução" in texto or "cumprimento de sentença" in texto:
        legendas.append("EMBARGOS À EXECUÇÃO")
    elif "distribuição" in texto:
        legendas.append("RELATÓRIO DISTRIBUIÇÃO")
        legendas.append("ABRIR FICHA NO PROMAD")
    else:
        legendas.append("ANÁLISE")

    return legendas


def extrair_texto(doc):
    texto = ""

    # Ler texto fora de tabelas
    for paragrafo in doc.paragraphs:
        texto += paragrafo.text + "\n"

    # Ler texto dentro de tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    texto += paragrafo.text + "\n"

    return texto


def separar_publicacoes(texto):
    return texto.split("COMENTÁRIOS")


def processar_documento(doc):
    texto_completo = extrair_texto(doc)
    publicacoes = separar_publicacoes(texto_completo)

    novo_doc = Document()

    for publicacao in publicacoes:
        if publicacao.strip() == "":
            continue

        # Adiciona o texto da publicação
        novo_doc.add_paragraph(publicacao.strip())

        # Gera legendas
        legendas = gerar_legendas(publicacao)

        for legenda in legendas:
            par = novo_doc.add_paragraph(legenda.upper())  # Legenda em CAIXA ALTA
            par.runs[0].bold = True
            aplicar_highlight(par)

        # Linha separadora
        novo_doc.add_paragraph("\n--------------------------------------------------\n")

    return novo_doc


# 🔥 INTERFACE STREAMLIT
st.title("🚀 Robô Shinji")

st.markdown("""
Envie seu documento (.docx) com as publicações. 
""")

arquivo = st.file_uploader("📤 Envie seu documento (.docx)", type="docx")

if arquivo is not None:
    doc = Document(arquivo)
    documento_processado = processar_documento(doc)

    buffer = BytesIO()
    documento_processado.save(buffer)
    buffer.seek(0)

    st.success("✅ Documento processado com sucesso!")
    st.download_button(
        label="📥 Baixar Documento com Legendas",
        data=buffer,
        file_name="documento_processado.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
